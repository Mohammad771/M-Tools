from flask import Flask, Blueprint, render_template, request, current_app
from docx2pdf import convert
import pythoncom
from PyPDF2 import PdfMerger
from docx import Document
from docx.shared import Inches, Pt, Cm


auto_certification_bp = Blueprint('auto_certification', __name__, url_prefix='/auto_certification')



@auto_certification_bp.route('/', methods=['GET', 'POST'])
def auto_certification():
    output_folder = current_app.config["OUTPUT_FOLDER"]
    upload_folder = current_app.config["UPLOAD_FOLDER"]
    # print("output folder: " + output_folder)
    if request.method == 'POST':

        uploaded_files = request.files.getlist('files')

        # Filter and process the uploaded files
        word_file = None
        image_files = []
        pdf_files = []

        word_file_path = ""
        pdf_file_path = ""

        for uploaded_file in uploaded_files:
            filename = uploaded_file.filename.lower()
            if filename.endswith('.docx'):
                word_file_path = upload_folder + uploaded_file.filename
                generated_pdf_file_path = current_app.config["OUTPUT_FOLDER"] + uploaded_file.filename.replace(".docx", ".pdf")
                uploaded_file.save(word_file_path)


            elif filename.endswith(('.jpg', '.jpeg', '.png')):
                image_files.append(uploaded_file)
                uploaded_file_path = upload_folder + uploaded_file.filename
                uploaded_file.save(uploaded_file_path)


            elif filename.endswith('.pdf'):
                pdf_files.append(uploaded_file)
                pdf_file_path = upload_folder + uploaded_file.filename
                uploaded_file.save(pdf_file_path)


        # add the stamp to the word file
        image_path = 'C:/Users/mohmd/OneDrive/Desktop/Programming/Projects/M_Tools/auto_certification/Stamp.png'
        add_stamp_image(word_file_path, image_path)
        # convert the word file into a pdf file
        convert_to_pdf(word_file_path)

        # compine files in a one pdf file
        files_paths = [generated_pdf_file_path, pdf_file_path]
        pdf_path = merge_pdfs(files_paths)

        return render_template('auto_certification/auto_certification.html', pdf_path=pdf_path, active_link='auto_certification')
    return render_template('auto_certification/auto_certification.html', active_link='auto_certification')    


"""
1. Receive files from the post array
2. Determine the word and the pdf file (or images)
3. convert the word to pdf, and convert the images to pdf (if any)
4. merge the first page, the word-converted pdf and the pdf or images files
5. return the path of the merged file to the user for download

"""
def convert_to_pdf(word_file_path):
    # print("word file received: " + word_file_path)
    pythoncom.CoInitialize()
    output_path = current_app.config["OUTPUT_FOLDER"]

    convert(word_file_path, output_path)
    return output_path


def merge_pdfs(input_files):

    print(input_files)

    output_file_path = current_app.config["OUTPUT_FOLDER"] + "Success.pdf"
    # Create a PdfMerger object
    merger = PdfMerger()

    # add the first page
    first_page_path = 'C:/Users/mohmd/OneDrive/Desktop/Programming/Projects/M_Tools/auto_certification/certification_page.pdf'
    merger.append(first_page_path)

    # Iterate over each input file
    for file in input_files:
        # Open the input PDF file in read-binary mode
        with open(file, 'rb') as f:
            # Add the PDF file to the merger
            merger.append(f)

    # Write the merged PDF to the output file
    with open(output_file_path, 'wb') as f:
        merger.write(f)

    return output_file_path


def add_stamp_image(word_file_path, image_path):

    print("I'm stamping")
    document = Document(word_file_path)

    # Add a new paragraph to the document
    paragraph = document.add_paragraph()

    # Load the image and add it to the paragraph
    run = paragraph.add_run()
    run.add_picture(image_path, width=Pt(166.5), height=Pt(109.5))

    paragraph.alignment = 1  # 1 stands for 'center'

    # Save the modified Word file
    document.save(word_file_path)

